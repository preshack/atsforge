from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Response, Request, Form
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import httpx
from io import BytesIO
import json
import re

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'atsforge_jwt_secret_key_2024')
JWT_ALGORITHM = "HS256"
OPENROUTER_KEY = os.environ.get('OPENROUTER_KEY')

# Use FREE DeepSeek model
AI_MODEL = "deepseek/deepseek-r1:free"

app = FastAPI(title="ATSForge API")
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ========================
# MODELS
# ========================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    onboarding_completed: bool = False
    created_at: datetime

class OnboardingData(BaseModel):
    experience_level: str
    industry: str
    target_roles: List[str]
    country: str
    resume_style: str

class ResumeCreate(BaseModel):
    title: str
    content: Optional[Dict[str, Any]] = None

class ResumeUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[Dict[str, Any]] = None

class Resume(BaseModel):
    resume_id: str
    user_id: str
    title: str
    content: Dict[str, Any]
    latex_code: Optional[str] = None
    ats_score: Optional[int] = None
    version: int = 1
    created_at: datetime
    updated_at: datetime

class ChatMessage(BaseModel):
    role: str
    content: str

class ResumeGenerateRequest(BaseModel):
    messages: List[ChatMessage]
    resume_id: Optional[str] = None
    section: Optional[str] = None

class LatexGenerateRequest(BaseModel):
    content: Dict[str, Any]
    template: str = "modern"

class CoverLetterCreate(BaseModel):
    title: str
    job_description: str
    resume_id: Optional[str] = None
    tone: str = "professional"
    company_name: Optional[str] = None
    position: Optional[str] = None

class CoverLetter(BaseModel):
    cover_letter_id: str
    user_id: str
    title: str
    content: str
    job_description: str
    tone: str
    created_at: datetime
    updated_at: datetime

class ATSOptimizeRequest(BaseModel):
    resume_content: Dict[str, Any]
    job_description: Optional[str] = None

class ATSScoreResponse(BaseModel):
    score: int
    issues: List[Dict[str, str]]
    suggestions: List[str]
    keyword_matches: List[str]
    missing_keywords: List[str]

# ========================
# HEALTH CHECK
# ========================

@api_router.get("/")
async def health_check():
    return {"message": "ATSForge API is running", "status": "healthy", "ai_model": AI_MODEL}

# ========================
# AUTH HELPERS
# ========================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str) -> str:
    payload = {
        "user_id": user_id,
        "exp": datetime.now(timezone.utc) + timedelta(days=7)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(request: Request) -> User:
    token = request.cookies.get("session_token")
    
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header.split(" ")[1]
    
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if session:
        expires_at = session.get("expires_at")
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at)
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        if expires_at < datetime.now(timezone.utc):
            raise HTTPException(status_code=401, detail="Session expired")
        
        user_doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
        if not user_doc:
            raise HTTPException(status_code=401, detail="User not found")
        return User(**user_doc)
    
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if not user_doc:
            raise HTTPException(status_code=401, detail="User not found")
        return User(**user_doc)
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ========================
# AI HELPER - FREE DeepSeek
# ========================

async def call_ai(system_prompt: str, user_prompt: str, max_tokens: int = 2000) -> str:
    """Call FREE DeepSeek model via OpenRouter"""
    try:
        async with httpx.AsyncClient(timeout=120.0) as http_client:
            response = await http_client.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENROUTER_KEY}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://atsforge.app",
                    "X-Title": "ATSForge"
                },
                json={
                    "model": AI_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "max_tokens": max_tokens,
                    "temperature": 0.7
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                content = result["choices"][0]["message"]["content"]
                # DeepSeek R1 may include <think> tags, strip them
                content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
                return content
            else:
                logger.error(f"OpenRouter error: {response.status_code} - {response.text}")
                return None
    except Exception as e:
        logger.error(f"AI call failed: {e}")
        return None

# ========================
# LaTeX TEMPLATES
# ========================

def generate_latex_resume(content: Dict[str, Any], template: str = "modern") -> str:
    """Generate LaTeX code for resume"""
    personal = content.get("personalInfo", {})
    
    # Escape LaTeX special characters
    def escape_latex(text):
        if not text:
            return ""
        special_chars = {'&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_', '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\textasciicircum{}'}
        for char, replacement in special_chars.items():
            text = text.replace(char, replacement)
        return text
    
    name = escape_latex(personal.get("name", "Your Name"))
    email = escape_latex(personal.get("email", "email@example.com"))
    phone = escape_latex(personal.get("phone", ""))
    location = escape_latex(personal.get("location", ""))
    linkedin = escape_latex(personal.get("linkedin", ""))
    summary = escape_latex(content.get("summary", ""))
    
    # Build experience section
    experience_tex = ""
    for exp in content.get("experience", []):
        title = escape_latex(exp.get("title", ""))
        company = escape_latex(exp.get("company", ""))
        dates = f"{escape_latex(exp.get('startDate', ''))} -- {escape_latex(exp.get('endDate', 'Present'))}"
        
        bullets_tex = ""
        for bullet in exp.get("bullets", []):
            if bullet:
                bullets_tex += f"    \\item {escape_latex(bullet)}\n"
        
        experience_tex += f"""
\\subsection*{{{title}}}
\\textit{{{company}}} \\hfill {dates}
\\begin{{itemize}}[leftmargin=*]
{bullets_tex}\\end{{itemize}}
"""

    # Build education section
    education_tex = ""
    for edu in content.get("education", []):
        degree = escape_latex(edu.get("degree", ""))
        field = escape_latex(edu.get("field", ""))
        school = escape_latex(edu.get("school", ""))
        grad_date = escape_latex(edu.get("graduationDate", ""))
        
        education_tex += f"""
\\textbf{{{degree} in {field}}} \\hfill {grad_date} \\\\
\\textit{{{school}}}
\\vspace{{0.5em}}
"""

    # Build skills
    skills = content.get("skills", [])
    skills_tex = ", ".join([escape_latex(s) for s in skills]) if skills else "Add your skills here"

    latex_code = f"""\\documentclass[11pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[T1]{{fontenc}}
\\usepackage{{lmodern}}
\\usepackage[margin=0.75in]{{geometry}}
\\usepackage{{enumitem}}
\\usepackage{{hyperref}}
\\usepackage{{xcolor}}
\\usepackage{{titlesec}}

% Colors
\\definecolor{{headercolor}}{{RGB}}{{79, 70, 229}}
\\definecolor{{linkcolor}}{{RGB}}{{79, 70, 229}}

% Section formatting
\\titleformat{{\\section}}{{\\Large\\bfseries\\color{{headercolor}}}}{{}}{{0em}}{{}}[\\titlerule]
\\titlespacing{{\\section}}{{0pt}}{{1em}}{{0.5em}}

% Hyperlink setup
\\hypersetup{{
    colorlinks=true,
    linkcolor=linkcolor,
    urlcolor=linkcolor
}}

% Remove page numbers
\\pagestyle{{empty}}

\\begin{{document}}

% Header
\\begin{{center}}
    {{\\Huge\\bfseries {name}}}\\\\[0.3em]
    {email}{f" $|$ {phone}" if phone else ""}{f" $|$ {location}" if location else ""}
    {f"\\\\\\href{{https://{linkedin}}}{{{linkedin}}}" if linkedin else ""}
\\end{{center}}

\\vspace{{0.5em}}

% Summary
{"\\section*{Professional Summary}" if summary else ""}
{summary if summary else ""}

% Experience
{"\\section*{Experience}" if experience_tex else ""}
{experience_tex}

% Education
{"\\section*{Education}" if education_tex else ""}
{education_tex}

% Skills
\\section*{{Skills}}
{skills_tex}

\\end{{document}}
"""
    return latex_code

# ========================
# AUTH ENDPOINTS
# ========================

@api_router.post("/auth/register")
async def register(user_data: UserCreate, response: Response):
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    hashed_pw = hash_password(user_data.password)
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "password_hash": hashed_pw,
        "picture": None,
        "onboarding_completed": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id)
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7*24*60*60
    )
    
    return {
        "user_id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "token": token,
        "onboarding_completed": False
    }

@api_router.post("/auth/login")
async def login(credentials: UserLogin, response: Response):
    user_doc = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(credentials.password, user_doc.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_token(user_doc["user_id"])
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7*24*60*60
    )
    
    return {
        "user_id": user_doc["user_id"],
        "email": user_doc["email"],
        "name": user_doc["name"],
        "picture": user_doc.get("picture"),
        "token": token,
        "onboarding_completed": user_doc.get("onboarding_completed", False)
    }

@api_router.post("/auth/session")
async def process_google_session(request: Request, response: Response):
    body = await request.json()
    session_id = body.get("session_id")
    
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id required")
    
    async with httpx.AsyncClient() as http_client:
        try:
            resp = await http_client.get(
                "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": session_id}
            )
            if resp.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid session")
            
            session_data = resp.json()
        except Exception as e:
            logger.error(f"Error fetching session data: {e}")
            raise HTTPException(status_code=500, detail="Auth service error")
    
    email = session_data.get("email")
    name = session_data.get("name")
    picture = session_data.get("picture")
    session_token = session_data.get("session_token")
    
    user_doc = await db.users.find_one({"email": email}, {"_id": 0})
    
    if user_doc:
        user_id = user_doc["user_id"]
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"name": name, "picture": picture}}
        )
        onboarding_completed = user_doc.get("onboarding_completed", False)
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        user_doc = {
            "user_id": user_id,
            "email": email,
            "name": name,
            "picture": picture,
            "onboarding_completed": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.users.insert_one(user_doc)
        onboarding_completed = False
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=7)).isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7*24*60*60
    )
    
    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "picture": picture,
        "onboarding_completed": onboarding_completed
    }

@api_router.get("/auth/me")
async def get_me(user: User = Depends(get_current_user)):
    return {
        "user_id": user.user_id,
        "email": user.email,
        "name": user.name,
        "picture": user.picture,
        "onboarding_completed": user.onboarding_completed
    }

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    token = request.cookies.get("session_token")
    if token:
        await db.user_sessions.delete_one({"session_token": token})
    
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out"}

# ========================
# ONBOARDING
# ========================

@api_router.post("/users/onboarding")
async def save_onboarding(data: OnboardingData, user: User = Depends(get_current_user)):
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {"onboarding": data.model_dump(), "onboarding_completed": True}}
    )
    return {"message": "Onboarding completed", "onboarding_completed": True}

@api_router.get("/users/onboarding")
async def get_onboarding(user: User = Depends(get_current_user)):
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    return {"onboarding": user_doc.get("onboarding"), "onboarding_completed": user_doc.get("onboarding_completed", False)}

# ========================
# RESUME ENDPOINTS
# ========================

@api_router.post("/resumes")
async def create_resume(resume_data: ResumeCreate, user: User = Depends(get_current_user)):
    resume_id = f"resume_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    content = resume_data.content or {
        "personalInfo": {"name": "", "email": "", "phone": "", "location": "", "linkedin": "", "website": ""},
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "certifications": [],
        "projects": [],
        "languages": []
    }
    
    latex_code = generate_latex_resume(content)
    
    resume_doc = {
        "resume_id": resume_id,
        "user_id": user.user_id,
        "title": resume_data.title,
        "content": content,
        "latex_code": latex_code,
        "ats_score": None,
        "version": 1,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(resume_doc)
    resume_doc.pop("_id", None)
    
    return {
        "resume_id": resume_id,
        "user_id": user.user_id,
        "title": resume_data.title,
        "content": content,
        "latex_code": latex_code,
        "ats_score": None,
        "version": 1,
        "created_at": now,
        "updated_at": now
    }

@api_router.get("/resumes")
async def get_resumes(user: User = Depends(get_current_user)):
    resumes = await db.resumes.find({"user_id": user.user_id}, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return resumes

@api_router.get("/resumes/{resume_id}")
async def get_resume(resume_id: str, user: User = Depends(get_current_user)):
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return resume

@api_router.put("/resumes/{resume_id}")
async def update_resume(resume_id: str, update_data: ResumeUpdate, user: User = Depends(get_current_user)):
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    update_fields = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if update_data.title:
        update_fields["title"] = update_data.title
    if update_data.content:
        update_fields["content"] = update_data.content
        update_fields["latex_code"] = generate_latex_resume(update_data.content)
    
    update_fields["version"] = resume.get("version", 1) + 1
    
    await db.resumes.update_one({"resume_id": resume_id}, {"$set": update_fields})
    
    updated = await db.resumes.find_one({"resume_id": resume_id}, {"_id": 0})
    return updated

@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, user: User = Depends(get_current_user)):
    result = await db.resumes.delete_one({"resume_id": resume_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {"message": "Resume deleted"}

@api_router.post("/resumes/{resume_id}/duplicate")
async def duplicate_resume(resume_id: str, user: User = Depends(get_current_user)):
    original = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not original:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    new_id = f"resume_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    new_resume = {
        **original,
        "resume_id": new_id,
        "title": f"{original['title']} (Copy)",
        "version": 1,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(new_resume)
    new_resume.pop("_id", None)
    
    return new_resume

# ========================
# LaTeX GENERATION
# ========================

@api_router.post("/resumes/generate-latex")
async def generate_latex(request: LatexGenerateRequest, user: User = Depends(get_current_user)):
    """Generate LaTeX code from resume content"""
    latex_code = generate_latex_resume(request.content, request.template)
    return {"latex_code": latex_code}

@api_router.post("/resumes/{resume_id}/latex")
async def get_resume_latex(resume_id: str, user: User = Depends(get_current_user)):
    """Get or regenerate LaTeX for a resume"""
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Regenerate LaTeX
    latex_code = generate_latex_resume(resume.get("content", {}))
    
    # Update in DB
    await db.resumes.update_one({"resume_id": resume_id}, {"$set": {"latex_code": latex_code}})
    
    return {"latex_code": latex_code, "resume_id": resume_id}

# ========================
# PDF UPLOAD & PARSING
# ========================

@api_router.post("/resumes/upload")
async def upload_resume(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = file.filename.lower().split('.')[-1]
    file_content = await file.read()
    
    extracted_text = ""
    
    try:
        if file_ext == 'pdf':
            import pdfplumber
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
        elif file_ext in ['docx', 'doc']:
            from docx import Document
            doc = Document(BytesIO(file_content))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        elif file_ext == 'txt':
            extracted_text = file_content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or TXT")
    except Exception as e:
        logger.error(f"File parsing error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")
    
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")
    
    parsed_content = parse_resume_text(extracted_text)
    latex_code = generate_latex_resume(parsed_content)
    
    resume_id = f"resume_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    resume_doc = {
        "resume_id": resume_id,
        "user_id": user.user_id,
        "title": file.filename.rsplit('.', 1)[0],
        "content": parsed_content,
        "latex_code": latex_code,
        "raw_text": extracted_text,
        "ats_score": None,
        "version": 1,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(resume_doc)
    resume_doc.pop("_id", None)
    
    return {
        "resume_id": resume_id,
        "title": resume_doc["title"],
        "content": parsed_content,
        "latex_code": latex_code,
        "extracted_text": extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text
    }

def parse_resume_text(text: str) -> Dict[str, Any]:
    """Parse raw resume text into structured format"""
    content = {
        "personalInfo": {"name": "", "email": "", "phone": "", "location": "", "linkedin": "", "website": ""},
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "certifications": [],
        "projects": [],
        "languages": []
    }
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        content["personalInfo"]["email"] = email_match.group()
    
    # Extract phone
    phone_match = re.search(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3,6}[-\s\.]?[0-9]{3,6}', text)
    if phone_match:
        content["personalInfo"]["phone"] = phone_match.group()
    
    # Extract LinkedIn
    linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    if linkedin_match:
        content["personalInfo"]["linkedin"] = linkedin_match.group()
    
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line and len(line) < 50 and not '@' in line and not any(c.isdigit() for c in line[:5]):
            content["personalInfo"]["name"] = line
            break
    
    # Extract skills
    skills_section = False
    for line in lines:
        line_lower = line.lower().strip()
        if 'skill' in line_lower or 'technologies' in line_lower or 'proficiencies' in line_lower:
            skills_section = True
            continue
        if skills_section and line.strip():
            if any(keyword in line_lower for keyword in ['experience', 'education', 'project', 'certification']):
                skills_section = False
            else:
                skills = re.split(r'[,|•·]', line)
                for skill in skills:
                    skill = skill.strip()
                    if skill and len(skill) < 30:
                        content["skills"].append(skill)
    
    return content

# ========================
# AI GENERATION
# ========================

@api_router.post("/resumes/generate")
async def generate_resume_content(request: ResumeGenerateRequest, user: User = Depends(get_current_user)):
    """Generate or improve resume content using FREE DeepSeek"""
    
    system_prompt = """You are an expert ATS-friendly resume writer. Help create professional resumes.

RULES:
1. Never fabricate experience or skills - only use what user provides
2. Use strong action verbs (Led, Developed, Implemented, Achieved, Increased, Reduced)
3. Include quantified achievements with numbers/percentages when possible
4. Keep formatting simple - no tables or graphics (ATS compatible)
5. Be concise - 1-2 lines per bullet point

When asked to generate full resume content, respond with JSON:
{
  "personalInfo": {"name": "", "email": "", "phone": "", "location": "", "linkedin": ""},
  "summary": "2-3 sentence professional summary",
  "experience": [{"company": "", "title": "", "location": "", "startDate": "", "endDate": "", "bullets": []}],
  "education": [{"school": "", "degree": "", "field": "", "graduationDate": ""}],
  "skills": ["skill1", "skill2"],
  "certifications": []
}

Otherwise, provide helpful resume writing advice and suggestions."""

    messages_text = "\n".join([f"{m.role}: {m.content}" for m in request.messages])
    
    ai_response = await call_ai(system_prompt, messages_text, 2000)
    
    if ai_response:
        return {"content": ai_response, "message": "Content generated successfully", "model": AI_MODEL}
    else:
        return {
            "content": "I can help you create an ATS-optimized resume! Please share:\n\n1. Your target job title\n2. Years of experience\n3. Key skills and achievements\n\nI'll help craft compelling bullet points with action verbs and metrics.",
            "message": "Using guided mode - AI temporarily unavailable",
            "model": "fallback"
        }

@api_router.post("/resumes/optimize")
async def optimize_resume(request: ATSOptimizeRequest, user: User = Depends(get_current_user)):
    """Analyze resume for ATS compatibility using FREE DeepSeek"""
    
    # Calculate local score first
    score_data = calculate_ats_score(request.resume_content, request.job_description)
    
    # Try AI for better suggestions
    system_prompt = """Analyze this resume for ATS compatibility. Return ONLY valid JSON:
{
  "score": 0-100,
  "issues": [{"severity": "high/medium/low", "issue": "description", "section": "section_name"}],
  "suggestions": ["suggestion1", "suggestion2"],
  "keyword_matches": ["keyword1"],
  "missing_keywords": ["keyword1"]
}"""
    
    resume_text = json.dumps(request.resume_content, indent=2)
    job_desc = request.job_description or "General professional position"
    
    ai_response = await call_ai(system_prompt, f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc}", 1500)
    
    if ai_response:
        try:
            json_match = re.search(r'\{[\s\S]*\}', ai_response)
            if json_match:
                analysis = json.loads(json_match.group())
                return ATSScoreResponse(
                    score=analysis.get("score", score_data["score"]),
                    issues=analysis.get("issues", score_data["issues"]),
                    suggestions=analysis.get("suggestions", score_data["suggestions"]),
                    keyword_matches=analysis.get("keyword_matches", score_data["keyword_matches"]),
                    missing_keywords=analysis.get("missing_keywords", score_data["missing_keywords"])
                )
        except:
            pass
    
    return ATSScoreResponse(**score_data)

def calculate_ats_score(content: Dict, job_description: str = None) -> Dict:
    """Calculate ATS score locally"""
    score = 50
    issues = []
    suggestions = []
    keyword_matches = []
    missing_keywords = []
    
    personal = content.get("personalInfo", {})
    
    if personal.get("name"):
        score += 5
    else:
        issues.append({"severity": "high", "issue": "Missing name", "section": "personalInfo"})
    
    if personal.get("email"):
        score += 5
    else:
        issues.append({"severity": "high", "issue": "Missing email", "section": "personalInfo"})
    
    if personal.get("phone"):
        score += 5
    else:
        issues.append({"severity": "medium", "issue": "Missing phone number", "section": "personalInfo"})
    
    if content.get("summary") and len(content["summary"]) > 50:
        score += 10
    else:
        issues.append({"severity": "medium", "issue": "Professional summary missing or too short", "section": "summary"})
        suggestions.append("Add a 2-3 sentence professional summary highlighting key achievements")
    
    experience = content.get("experience", [])
    if len(experience) > 0:
        score += 10
        for exp in experience:
            if exp.get("bullets") and len(exp["bullets"]) > 0:
                score += 2
                for bullet in exp["bullets"]:
                    if bullet and any(verb in bullet.lower() for verb in ["led", "developed", "implemented", "achieved", "increased", "decreased", "managed", "created"]):
                        score += 1
                        break
    else:
        issues.append({"severity": "high", "issue": "No work experience listed", "section": "experience"})
    
    skills = content.get("skills", [])
    if len(skills) >= 5:
        score += 10
        keyword_matches.extend(skills[:5])
    elif len(skills) > 0:
        score += 5
        keyword_matches.extend(skills)
    else:
        issues.append({"severity": "medium", "issue": "No skills listed", "section": "skills"})
        suggestions.append("Add 8-12 relevant technical and soft skills")
    
    if content.get("education") and len(content["education"]) > 0:
        score += 5
    
    if job_description:
        job_words = set(re.findall(r'\b\w+\b', job_description.lower()))
        resume_text = json.dumps(content).lower()
        resume_words = set(re.findall(r'\b\w+\b', resume_text))
        
        stop_words = {"the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "is", "are", "was", "were", "be", "been", "have", "has", "had", "do", "does", "did", "will", "would", "could", "should", "may", "might", "must", "can", "need"}
        job_keywords = job_words - stop_words
        
        matches = job_keywords & resume_words
        missing = job_keywords - resume_words
        
        keyword_matches.extend(list(matches)[:10])
        missing_keywords.extend(list(missing)[:10])
        
        match_ratio = len(matches) / len(job_keywords) if job_keywords else 0
        score += int(match_ratio * 20)
    
    score = min(score, 100)
    
    if not suggestions:
        suggestions = [
            "Use strong action verbs at the start of bullet points",
            "Quantify achievements with numbers and percentages",
            "Tailor your resume to each job application"
        ]
    
    return {
        "score": score,
        "issues": issues,
        "suggestions": suggestions,
        "keyword_matches": keyword_matches,
        "missing_keywords": missing_keywords
    }

# ========================
# COVER LETTER
# ========================

@api_router.post("/cover-letters")
async def create_cover_letter(data: CoverLetterCreate, user: User = Depends(get_current_user)):
    """Generate cover letter using FREE DeepSeek"""
    
    resume_content = ""
    if data.resume_id:
        resume = await db.resumes.find_one({"resume_id": data.resume_id, "user_id": user.user_id}, {"_id": 0})
        if resume:
            resume_content = json.dumps(resume.get("content", {}), indent=2)
    
    tone_map = {
        "professional": "formal, professional tone suitable for corporate environments",
        "friendly": "warm, approachable yet professional tone",
        "confident": "confident, assertive tone highlighting achievements",
        "enthusiastic": "enthusiastic, energetic tone showing passion"
    }
    
    system_prompt = f"""Write a compelling cover letter with a {tone_map.get(data.tone, 'professional tone')}.

RULES:
1. 3-4 paragraphs maximum
2. Address specific requirements from the job description
3. Avoid generic phrases like "I am writing to apply"
4. Strong opening hook and clear call to action
5. Professional closing

Write the cover letter directly without any preamble."""

    user_prompt = f"""Company: {data.company_name or 'the company'}
Position: {data.position or 'the position'}

Job Description:
{data.job_description}

{"Resume/Background:\n" + resume_content if resume_content else ""}

Write the cover letter:"""

    ai_response = await call_ai(system_prompt, user_prompt, 1500)
    
    if not ai_response:
        ai_response = f"""Dear Hiring Manager,

I am excited to apply for the {data.position or 'position'} at {data.company_name or 'your company'}. With my relevant background and skills, I am confident I can make meaningful contributions to your team.

Your job description particularly resonates with my experience and career aspirations. I have developed expertise that directly aligns with your requirements and am eager to bring this value to your organization.

I would welcome the opportunity to discuss how my skills and experience could benefit {data.company_name or 'your organization'}. Thank you for considering my application, and I look forward to hearing from you.

Best regards"""

    cover_letter_id = f"cl_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    cl_doc = {
        "cover_letter_id": cover_letter_id,
        "user_id": user.user_id,
        "title": data.title,
        "content": ai_response,
        "job_description": data.job_description,
        "tone": data.tone,
        "created_at": now,
        "updated_at": now
    }
    
    await db.cover_letters.insert_one(cl_doc)
    cl_doc.pop("_id", None)
    
    return cl_doc

@api_router.get("/cover-letters")
async def get_cover_letters(user: User = Depends(get_current_user)):
    letters = await db.cover_letters.find({"user_id": user.user_id}, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return letters

@api_router.get("/cover-letters/{cover_letter_id}")
async def get_cover_letter(cover_letter_id: str, user: User = Depends(get_current_user)):
    cl = await db.cover_letters.find_one({"cover_letter_id": cover_letter_id, "user_id": user.user_id}, {"_id": 0})
    if not cl:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return cl

@api_router.put("/cover-letters/{cover_letter_id}")
async def update_cover_letter(cover_letter_id: str, request: Request, user: User = Depends(get_current_user)):
    body = await request.json()
    content = body.get("content", "")
    
    result = await db.cover_letters.update_one(
        {"cover_letter_id": cover_letter_id, "user_id": user.user_id},
        {"$set": {"content": content, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return {"message": "Updated"}

@api_router.delete("/cover-letters/{cover_letter_id}")
async def delete_cover_letter(cover_letter_id: str, user: User = Depends(get_current_user)):
    result = await db.cover_letters.delete_one({"cover_letter_id": cover_letter_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    return {"message": "Deleted"}

# ========================
# EXPORT
# ========================

@api_router.post("/export/pdf/{resume_id}")
async def export_resume_pdf(resume_id: str, user: User = Depends(get_current_user)):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    content = resume.get("content", {})
    buffer = BytesIO()
    
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.75*inch, leftMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, spaceAfter=4, alignment=1)
    contact_style = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=9, alignment=1, textColor=HexColor('#666666'))
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=11, spaceBefore=10, spaceAfter=4, textColor=HexColor('#4f46e5'))
    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10, spaceAfter=3)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=10, leftIndent=15, spaceAfter=2)
    
    story = []
    
    personal = content.get("personalInfo", {})
    if personal.get("name"):
        story.append(Paragraph(personal["name"], title_style))
    
    contact_parts = [p for p in [personal.get("email"), personal.get("phone"), personal.get("location"), personal.get("linkedin")] if p]
    if contact_parts:
        story.append(Paragraph(" • ".join(contact_parts), contact_style))
    
    story.append(Spacer(1, 8))
    
    if content.get("summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(content["summary"], normal_style))
    
    if content.get("experience"):
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in content["experience"]:
            story.append(Paragraph(f"<b>{exp.get('title', '')}</b> | {exp.get('company', '')} | {exp.get('startDate', '')} - {exp.get('endDate', 'Present')}", normal_style))
            for bullet in exp.get("bullets", []):
                if bullet:
                    story.append(Paragraph(f"• {bullet}", bullet_style))
    
    if content.get("education"):
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in content["education"]:
            story.append(Paragraph(f"<b>{edu.get('degree', '')} in {edu.get('field', '')}</b> | {edu.get('school', '')} | {edu.get('graduationDate', '')}", normal_style))
    
    if content.get("skills"):
        story.append(Paragraph("SKILLS", heading_style))
        story.append(Paragraph(", ".join(content["skills"]), normal_style))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={resume['title']}.pdf"})

@api_router.post("/export/docx/{resume_id}")
async def export_resume_docx(resume_id: str, user: User = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    content = resume.get("content", {})
    doc = Document()
    
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    
    personal = content.get("personalInfo", {})
    if personal.get("name"):
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal["name"])
        name_run.font.size = Pt(16)
        name_run.bold = True
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    contact_parts = [p for p in [personal.get("email"), personal.get("phone"), personal.get("location")] if p]
    if contact_parts:
        contact_para = doc.add_paragraph(" • ".join(contact_parts))
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if content.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(content["summary"])
    
    if content.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in content["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')}").bold = True
            p.add_run(f" | {exp.get('company', '')} | {exp.get('startDate', '')} - {exp.get('endDate', 'Present')}")
            for bullet in exp.get("bullets", []):
                if bullet:
                    doc.add_paragraph(f"• {bullet}")
    
    if content.get("education"):
        doc.add_heading("Education", level=1)
        for edu in content["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
            doc.add_paragraph(f"{edu.get('school', '')} | {edu.get('graduationDate', '')}")
    
    if content.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(content["skills"]))
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={resume['title']}.docx"})

@api_router.get("/export/latex/{resume_id}")
async def export_resume_latex(resume_id: str, user: User = Depends(get_current_user)):
    """Export resume as LaTeX source code"""
    resume = await db.resumes.find_one({"resume_id": resume_id, "user_id": user.user_id}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    latex_code = resume.get("latex_code") or generate_latex_resume(resume.get("content", {}))
    
    return Response(
        content=latex_code,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={resume['title']}.tex"}
    )

@api_router.post("/export/cover-letter/pdf/{cover_letter_id}")
async def export_cover_letter_pdf(cover_letter_id: str, user: User = Depends(get_current_user)):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    
    cl = await db.cover_letters.find_one({"cover_letter_id": cover_letter_id, "user_id": user.user_id}, {"_id": 0})
    if not cl:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    paragraphs = cl["content"].split("\n\n")
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.replace("\n", "<br/>"), styles["Normal"]))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={cl['title']}.pdf"})

# ========================
# DASHBOARD
# ========================

@api_router.get("/dashboard/stats")
async def get_dashboard_stats(user: User = Depends(get_current_user)):
    resume_count = await db.resumes.count_documents({"user_id": user.user_id})
    cover_letter_count = await db.cover_letters.count_documents({"user_id": user.user_id})
    
    recent_resumes = await db.resumes.find(
        {"user_id": user.user_id}, 
        {"_id": 0, "resume_id": 1, "title": 1, "updated_at": 1, "ats_score": 1}
    ).sort("updated_at", -1).limit(5).to_list(5)
    
    recent_cover_letters = await db.cover_letters.find(
        {"user_id": user.user_id},
        {"_id": 0, "cover_letter_id": 1, "title": 1, "updated_at": 1}
    ).sort("updated_at", -1).limit(5).to_list(5)
    
    avg_score = None
    scored_resumes = await db.resumes.find(
        {"user_id": user.user_id, "ats_score": {"$ne": None}},
        {"_id": 0, "ats_score": 1}
    ).to_list(100)
    if scored_resumes:
        avg_score = sum(r["ats_score"] for r in scored_resumes) // len(scored_resumes)
    
    return {
        "resume_count": resume_count,
        "cover_letter_count": cover_letter_count,
        "recent_resumes": recent_resumes,
        "recent_cover_letters": recent_cover_letters,
        "avg_ats_score": avg_score
    }

@api_router.get("/templates")
async def get_templates():
    return [
        {"id": "modern", "name": "Modern Professional", "description": "Clean, minimal design"},
        {"id": "classic", "name": "Classic Executive", "description": "Traditional corporate format"},
        {"id": "creative", "name": "Creative Bold", "description": "Stand-out design"},
        {"id": "technical", "name": "Technical Focus", "description": "Skills-first layout"}
    ]

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
